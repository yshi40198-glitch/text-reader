# -*- coding: utf-8 -*-
"""文档内容提取：PDF / Word(.docx .doc) / 纯文本 / EPUB / MOBI / FB2，含朗读整理工具"""
import os
import re


_BOX_CHARS = set(chr(c) for c in range(0x2500, 0x2580))  # 表格框线等装饰字符

_PAGE_ONLY_RE = re.compile(r'^[-–—·\s]*\d{1,4}[-–—·\s]*$')       # 纯页码
_PAGE_LABEL_RE = re.compile(r'^第\s*\d{1,4}\s*[页章]\s*$')        # 第X页/第X章
_PAGE_DASH_RE = re.compile(r'^\s*[-–—]\s*\d{1,4}\s*[-–—]\s*$')   # - 12 -
_FOOTER_RE = re.compile(r'^(页码|第\s*\d+\s*页\s*共\s*\d+\s*页|'
                        r'Copyright|©|All\s+Rights\s+Reserved)',
                        re.I)


def _is_noise(line):
    """智能过滤：判断一行是否为页码/页眉页脚等朗读噪音。"""
    if not line:
        return False
    if _PAGE_ONLY_RE.match(line) and len(line) <= 12:
        return True
    if _PAGE_LABEL_RE.match(line):
        return True
    if _PAGE_DASH_RE.match(line):
        return True
    if _FOOTER_RE.match(line) and len(line) <= 30:
        return True
    return False


def clean_for_speech(text, smart_filter=True):
    """整理成适合朗读的文本：去掉表格框线/装饰字符、压缩连续空行；
    smart_filter=True 时跳过页码、页眉页脚等噪音行。"""
    lines = []
    for line in (text or '').splitlines():
        line = ''.join(ch for ch in line if ch not in _BOX_CHARS).strip()
        if line:
            if smart_filter and _is_noise(line):
                continue
            lines.append(line)
        else:
            if lines and lines[-1] != '':
                lines.append('')
    out, blank = [], False
    for line in lines:
        if line == '':
            if not blank:
                out.append('')
            blank = True
        else:
            blank = False
            out.append(line)
    while out and out[-1] == '':
        out.pop()
    return '\n'.join(out)


def split_chunks(text, max_len=180):
    """按句切块，每块约 max_len 字；无标点的超长段也强制切分。"""
    import re
    sentences = re.split(r'(?<=[。！？；…\n])', text or '')
    out, buf = [], ''
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        while len(s) > max_len:
            out.append(s[:max_len])
            s = s[max_len:]
        if s:
            buf += s
        if len(buf) >= max_len:
            out.append(buf)
            buf = ''
    if buf:
        out.append(buf)
    return out


def extract_text(path):
    """按扩展名提取文本。读不了时抛异常(由上层转成友好提示)。"""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        return extract_pdf(path)
    elif ext == '.docx':
        return extract_docx(path)
    elif ext == '.doc':
        return extract_doc(path)
    elif ext in ('.epub',):
        return extract_epub(path)
    elif ext in ('.mobi', '.azw', '.azw3'):
        return extract_mobi(path)
    elif ext in ('.fb2',):
        return extract_fb2(path)
    elif ext in ('.html', '.htm'):
        return extract_html(path)
    elif ext in ('.txt', '.md', '.text'):
        return read_text_file(path)
    else:
        raise ValueError(f"暂不支持此格式 ({ext or '无扩展名'})")


def read_text_file(path):
    """读纯文本：自动识别常见中文编码(UTF-8 / GBK / UTF-16)，避免乱码。"""
    with open(path, 'rb') as f:
        raw = f.read()
    # 1) 尝试 UTF-8
    try:
        return raw.decode('utf-8')
    except (UnicodeDecodeError, ValueError):
        pass
    # 2) 尝试 GBK（中文 Windows ANSI 最常见）
    try:
        return raw.decode('gbk')
    except (UnicodeDecodeError, ValueError):
        pass
    # 3) 尝试 UTF-16 (带 BOM)
    for enc in ('utf-16', 'gb18030', 'big5'):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, ValueError):
            continue
    # 4) 最后兜底：忽略错误
    return raw.decode('utf-8', errors='ignore')


def extract_pdf(path):
    import pymupdf  # PyMuPDF
    doc = pymupdf.open(path)
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return '\n'.join(pages)


def extract_docx(path):
    """Word .docx（现代 XML 格式）用 python-docx 提取。"""
    import docx
    d = docx.Document(path)
    parts = []
    # 正文段落
    for para in d.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(' | '.join(cells))
    return '\n'.join(parts)


def extract_epub(path):
    """EPUB 电子书：本质是 zip，把里面的 XHTML/HTML 正文提取出来。"""
    import zipfile
    from html.parser import HTMLParser

    class _Text(HTMLParser):
        def __init__(self):
            super().__init__()
            self.parts = []

        def handle_data(self, data):
            self.parts.append(data)

    texts = []
    with zipfile.ZipFile(path) as z:
        names = [n for n in z.namelist()
                 if n.lower().endswith(('.xhtml', '.html', '.htm'))]
        names.sort()
        for n in names:
            try:
                p = _Text()
                p.feed(z.read(n).decode('utf-8', errors='ignore'))
                chunk = ''.join(p.parts).strip()
                if chunk:
                    texts.append(chunk)
            except Exception:
                continue
    return '\n'.join(texts)


def _mobi_unpack_palmdoc(data):
    """MOBI 用的 PalmDoc 压缩解压（LZ77 变体），返回解压后的字节。"""
    out = bytearray()
    p = 0
    n = len(data)
    while p < n:
        c = data[p]
        p += 1
        if 1 <= c <= 8:
            # 短字面量：直接复制 c 个字节
            out.extend(data[p:p + c])
            p += c
        elif c < 128:
            # 单个字面字节
            out.append(c)
        elif c >= 192:
            # 空格 + 一个字节（c ^ 128）
            out.append(0x20)
            out.append(c ^ 128)
        else:
            # 128..191：回指复制
            if p >= n:
                break
            c = (c << 8) | data[p]
            p += 1
            dist = (c >> 3) & 0x7FF
            length = (c & 7) + 3
            if dist > length:
                out.extend(out[-dist:length - dist])
            else:
                for _ in range(length):
                    out.append(out[-dist])
    return bytes(out)


def _mobi_strip_html(html_text):
    """把 MOBI 的 HTML 正文转成纯文本（解析实体，保留段落换行）。"""
    from html.parser import HTMLParser
    import html as _html

    class _Text(HTMLParser):
        def __init__(self):
            super().__init__()
            self.parts = []

        def handle_data(self, data):
            self.parts.append(data)

        def handle_entityref(self, name):
            self.parts.append('&%s;' % name)

        def handle_charref(self, name):
            self.parts.append('&#%s;' % name)

    p = _Text()
    p.feed(html_text)
    return _html.unescape(''.join(p.parts))


def extract_mobi(path):
    """MOBI / AZW 电子书：解析 PDB 结构 + PalmDoc 解压，取出正文。

    MOBI 本质是 Palm 数据库（PDB）：第 0 条记录是 PalmDoc 头 + MOBI 头，
    之后的正文记录用 PalmDoc 压缩算法存 HTML。这里只做纯文本提取，
    不解析目录、封面等元数据。
    """
    import struct

    with open(path, 'rb') as f:
        raw = f.read()
    if len(raw) < 78:
        raise ValueError('文件太小，不是有效的 MOBI 电子书')
    nrec = struct.unpack('>H', raw[76:78])[0]
    if nrec < 2:
        raise ValueError('MOBI 电子书里没有正文记录')
    offsets = []
    for i in range(nrec):
        offsets.append(struct.unpack('>I', raw[78 + i * 8: 78 + i * 8 + 4])[0])

    rec0 = raw[offsets[0]: offsets[1] if nrec > 1 else len(raw)]
    compression = struct.unpack('>H', rec0[0:2])[0]
    text_len = struct.unpack('>I', rec0[4:8])[0]

    # MOBI 头在第 0 条记录的 16 字节之后；第一个图片记录号在 0x5C 处
    first_image = None
    if len(rec0) >= 16 + 0x60 and rec0[16:20] == b'MOBI':
        first_image = struct.unpack('>I', rec0[16 + 0x5C: 16 + 0x5C + 4])[0]

    end = nrec
    if first_image and 1 < first_image < nrec:
        end = first_image

    parts = []
    for i in range(1, end):
        s = offsets[i]
        e = offsets[i + 1] if i + 1 < nrec else len(raw)
        seg = raw[s:e]
        if compression == 2:
            seg = _mobi_unpack_palmdoc(seg)
        parts.append(seg)
    data = b''.join(parts)
    if text_len and 0 < text_len < len(data):
        data = data[:text_len]

    text = None
    for enc in ('utf-8', 'gb18030'):
        try:
            text = data.decode(enc)
            break
        except (UnicodeDecodeError, ValueError):
            continue
    if text is None:
        text = data.decode('utf-8', errors='ignore')
    return _mobi_strip_html(text)


def extract_fb2(path):
    """FB2 电子书（XML）：把正文段落按顺序拼出来。"""
    import xml.etree.ElementTree as ET

    raw = read_text_file(path)
    root = ET.fromstring(raw)

    def local(tag):
        return tag.rsplit('}', 1)[-1]

    leaf_tags = ('p', 'subtitle', 'text-author', 'v')

    parts = []
    for el in root.iter():
        t = local(el.tag)
        if t not in leaf_tags:
            continue
        # 只取最底层的段落，避免父节点(如 title/poem)重复收集
        has_leaf_child = False
        for child in el.iter():
            if child is el:
                continue
            if local(child.tag) in leaf_tags:
                has_leaf_child = True
                break
        if not has_leaf_child:
            txt = ''.join(el.itertext()).strip()
            if txt:
                parts.append(txt)
    if not parts:
        raise ValueError('FB2 电子书里没有提取到正文')
    return '\n'.join(parts)


def extract_html(path):
    """网页/HTML 文件：去掉标签只留文字。"""
    from html.parser import HTMLParser

    class _Text(HTMLParser):
        def __init__(self):
            super().__init__()
            self.parts = []

        def handle_data(self, data):
            self.parts.append(data)

    p = _Text()
    p.feed(read_text_file(path))
    return ''.join(p.parts)


def extract_doc(path):
    """Word .doc（旧版二进制格式）读取。

    优先用本机安装的 Microsoft Word 通过 COM 打开并取文本（最可靠，
    需装了 Microsoft Office/Word）。若本机没装 Word，返回 None，
    由上层引导把文件另存为 .docx。
    """
    try:
        # 尝试用 Word 的 COM 接口打开（仅 Windows + 装了 Office 时有效）
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(path), ReadOnly=True)
            try:
                text = doc.Content.Text
            finally:
                doc.Close(False)
            return text
        finally:
            word.Quit()
    except Exception:
        pass
    return None


def extract_chunks(text, words_per_chunk=120):
    """按句切成长度不超 words_per_chunk 的分块，用于逐句朗读。"""
    import re
    sentences = re.split(r'(?<=[。！？；…\n.])\s*', text)
    chunks, cur, cnt = [], [], 0
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        cur.append(s)
        cnt += len(s)
        if cnt >= words_per_chunk:
            chunks.append(''.join(cur))
            cur, cnt = [], 0
    if cur:
        chunks.append(''.join(cur))
    return chunks
